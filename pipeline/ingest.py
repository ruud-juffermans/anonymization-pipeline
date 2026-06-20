"""Extract plain text from supported document formats.

Supported: .txt/.md/.csv, .docx, .xlsx, digital .pdf, .eml. Heavy parsers are
imported lazily so a missing optional dependency only fails for that format.
"""

from __future__ import annotations

from email import policy
from email.parser import BytesParser
from pathlib import Path


def extract_text(path: str | Path) -> str:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in {".txt", ".md", ".csv", ".log"}:
        return p.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        return _docx(p)
    if suffix == ".xlsx":
        return _xlsx(p)
    if suffix == ".pdf":
        return _pdf(p)
    if suffix in {".eml", ".msg"}:
        return _email(p)
    raise ValueError(f"Unsupported format: {suffix} ({p.name})")


def _docx(p: Path) -> str:
    from docx import Document
    doc = Document(str(p))
    parts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _xlsx(p: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(p), read_only=True, data_only=True)
    lines = []
    for ws in wb.worksheets:
        lines.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            lines.append("\t".join(cells))
    return "\n".join(lines)


def _pdf(p: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(p))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _email(p: Path) -> str:
    with open(p, "rb") as fh:
        msg = BytesParser(policy=policy.default).parse(fh)
    headers = [f"{h}: {msg[h]}" for h in ("From", "To", "Cc", "Subject", "Date")
               if msg[h]]
    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                body += part.get_content()
    else:
        body = msg.get_content()
    return "\n".join(headers) + "\n\n" + body
